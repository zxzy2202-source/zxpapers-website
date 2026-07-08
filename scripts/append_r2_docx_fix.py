from pathlib import Path
from docx import Document

base = Path.cwd()
path = next(base.glob('B2B*.docx'))
doc = Document(str(path))

start_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text or ''
    if ('634cec8' in text) or ('???R2' in text) or ('R2 ?????????' in text):
        start_idx = i
        break

if start_idx is not None:
    # Move to the heading paragraph if the match was inside the appended section.
    while start_idx > 0 and 'R2' not in (doc.paragraphs[start_idx].text or ''):
        start_idx -= 1
    for p in doc.paragraphs[start_idx:]:
        element = p._element
        element.getparent().remove(element)

marker = '二十、R2 图片公网地址与代理失效复盘'
doc.add_paragraph('')
doc.add_heading(marker, level=1)
doc.add_paragraph(
    '本节来自 zxpapers 网站 2026-06-29 线上图片故障复盘。问题表现为：前台页面可打开，'
    '但大量产品图、Hero 图不显示；HTML 中图片地址输出为站内代理路径 '
    'https://www.zxpapers.com/r2-assets/...，这些地址返回 404；同一路径的真实 Cloudflare R2 '
    '公网直链 https://pub-529e97a14b4f4353b8b72301cfd8b481.r2.dev/... 返回 200。'
)

doc.add_heading('20.1 故障根因', level=2)
for text in [
    '线上 NEXT_PUBLIC_R2_URL 被配置为或等效解析为站内 /r2-assets 代理地址，而该代理路径在当前部署中返回 404。',
    '数据库 / KV 中历史图片值可能已经保存为 /r2-assets/... 或 https://www.zxpapers.com/r2-assets/...；如果 resolver 直接信任该值，前台会继续输出坏链。',
    'Vercel / CDN / Next.js rewrite 不能作为图片公网源的唯一依赖；B2B 获客站图片必须优先输出真实可访问的 R2/CDN 公网地址。',
]:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('20.2 修复标准', level=2)
table = doc.add_table(rows=1, cols=3)
hdr = table.rows[0].cells
hdr[0].text = '位置'
hdr[1].text = '规则'
hdr[2].text = '验收'
rows = [
    ('src/lib/r2.ts', '新增 normalizeR2PublicBase。若 NEXT_PUBLIC_R2_URL 包含 /r2-assets，则视为误配置并回退到真实 R2 公网域名。', '前台 HTML 不再输出 www.zxpapers.com/r2-assets 图片地址。'),
    ('src/lib/r2.ts / r2Image()', '历史 /r2-assets/... 值必须强制转换为 R2_PUBLIC_BASE + path。', '旧 KV 图片槽位无需手工清库也能恢复显示。'),
    ('next.config.ts', 'getAbsolutePublicUrl 同样拒绝包含 /r2-assets 的配置，避免 rewrite/remotePatterns 指向站内代理。', '构建后 rewrite 不会自引用或继续指向失效代理。'),
    ('后台首页', '系统配置状态卡片必须显示 KV、R2 上传凭据、R2 Bucket、公网图片域名是否完整。', '后台能快速判断图片和保存问题源头。'),
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value

doc.add_heading('20.3 验收命令与判断口径', level=2)
for text in [
    '抓取首页和产品页 HTML，确认图片 URL 不再是 https://www.zxpapers.com/r2-assets/...。',
    '对代表性图片执行 curl -I：站内 /r2-assets 若返回 404，不影响最终验收；真实 R2 直链必须返回 200。',
    '执行 tsc --noEmit 和 next build，确认类型检查与生产构建通过。',
    '部署后清 CDN / 浏览器缓存，重新打开首页、产品页、OEM 页和图片管理相关页面。',
    '后台首页查看“系统配置状态”：KV 数据存储、R2 上传凭据、R2 Bucket、公网图片域名均应为正常或有明确提示。',
]:
    doc.add_paragraph(text, style='List Number')

doc.add_heading('20.4 避坑规则', level=2)
for text in [
    'NEXT_PUBLIC_R2_URL 只能填写真实公网图片域名，如 https://pub-xxxx.r2.dev 或绑定后的 CDN 域名，不得填写本站 /r2-assets 代理路径。',
    '图片槽位 resolver 必须集中处理历史 URL、相对路径、绝对 R2 地址和外部图片，不允许在业务组件里散落 URL 拼接逻辑。',
    '上传端返回的 url 与前台展示端使用同一 R2_PUBLIC_BASE，避免上传成功但展示地址分裂。',
    '出现图片异常时，先检查 HTML 实际输出 URL 和 HTTP 状态码，不要只看后台是否有图片记录。',
    '后台状态卡片只能显示配置是否完整和公网 base，不得输出 R2_SECRET_ACCESS_KEY、R2_ACCESS_KEY_ID、KV_REST_API_TOKEN 等敏感值。',
]:
    doc.add_paragraph(text, style='List Bullet')

doc.add_paragraph('本次 zxpapers 修复提交：634cec8 fix: normalize R2 public image base。该经验应纳入所有 H 轨 B2B 英文独立站图片系统验收基线。')

doc.save(str(path))
print(f'updated {path.name}')
